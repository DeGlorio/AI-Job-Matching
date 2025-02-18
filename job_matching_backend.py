import logging
import traceback
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from typing import List, Dict
import openai
from openai import OpenAIError  # ✅ Import OpenAI error handling correctly
import pdfplumber
import requests
import os
import io
import docx
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.responses import FileResponse
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
#from flask import Flask, request, send_file, jsonify - removed as we're using FastAPI rather than flask
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime, timedelta
from difflib import SequenceMatcher

# OpenAI API Key
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Initialize OpenAI client
client = openai.OpenAI(api_key=OPENAI_API_KEY)

# Initialize FastAPI app
app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://animated-meme-5grjvwvqr6jjh4q5g-3000.app.github.dev",
        "https://animated-meme-5grjvwvqr6jjh4q5g-8000.app.github.dev",
        "http://localhost:3000",
        "http://127.0.0.1:3000"
    ],  
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],  # Allow browser to read response headers
)
#add middleware to explicitly set headers
class CORSMiddlewareCustom(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["Access-Control-Allow-Origin"] = "https://animated-meme-5grjvwvqr6jjh4q5g-3000.app.github.dev"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        response.headers["Access-Control-Allow-Headers"] = "*"
        return response

app.add_middleware(CORSMiddlewareCustom)
# Storage for Job Seekers & Jobs
job_seekers = []
employer_jobs = []

logging.basicConfig(level=logging.INFO)

# Helper Function: Extract Text from Resume

def extract_resume_text(file: UploadFile):
    file_extension = file.filename.split(".")[-1].lower()
    file_content = file.file.read()  # Read file once
    file.file.seek(0)  # Reset cursor for re-reading

    try:
        if file_extension == "pdf":
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif file_extension == "docx":
            doc = docx.Document(io.BytesIO(file_content))
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

# ✅ Employers Can Post Jobs
class JobPost(BaseModel):
    job_title: str
    skills_required: List[str]
    questions: List[str]

@app.post("/add_job")
async def add_job(job: JobPost):
    employer_jobs.append({"title": job.job_title, "skills": job.skills_required, "questions": job.questions})
    return {"job": job}

# ✅ Get Job Postings
@app.get("/employer_jobs")
async def get_employer_jobs():
    return {"jobs": employer_jobs}

# ✅ Job Seekers Can Register
class JobSeeker(BaseModel):
    name: str
    email: str
    resume_text: str

@app.post("/register_job_seeker")
async def register_job_seeker(seeker: JobSeeker):
    job_seekers.append(seeker.dict())
    return {"message": "Job seeker profile stored successfully."}

# ✅ Resume Enhancement
@app.post("/enhance_resume")
async def enhance_resume(resume: UploadFile = File(...), job_description: str = Form(...)):
    if not resume:
        raise HTTPException(status_code=400, detail="No resume file uploaded")

    # Extract resume text
    try:
        resume_text = extract_resume_text(resume)
        logging.info(f"Extracted resume text: {resume_text[:500]}")  # Log first 500 chars
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Extracted resume text is empty.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting resume text: {str(e)}")

    # Ensure OpenAI API Key is set
    if not OPENAI_API_KEY:
        return JSONResponse(status_code=400, content={"error": "OpenAI API key is missing."})

    try:
        openai.api_key = OPENAI_API_KEY  # ✅ Set API Key
        logging.info("OpenAI API Key is set.")
        # **Generate AI-enhanced resume**
        resume_prompt = f"""
        Improve this resume to match the following job description.
        Resume: {resume_text}
        Job Description: {job_description}
        Provide a well-structured, professional version.
        """

        response_resume = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a career expert enhancing resumes."},
                {"role": "user", "content": resume_prompt}
            ]
        )

        enhanced_resume = response_resume.choices[0].message.content.strip()

        # **Generate AI-enhanced cover letter**
        cover_letter_prompt = f"""
        Write a personalized cover letter using the resume below for this job.
        Resume: {enhanced_resume}
        Job Description: {job_description}
        Make it professional, engaging, and concise.
        """

        response_cover_letter = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a professional cover letter writer."},
                {"role": "user", "content": cover_letter_prompt}
            ]
        )

        cover_letter = response_cover_letter.choices[0].message.content.strip()

        # Ensure responses are not empty
        if not enhanced_resume or not cover_letter:
            raise HTTPException(status_code=500, detail="AI response was empty. Please try again.")

        return {
            "resume_text": resume_text,  # ✅ Ensure resume_text is included
            "enhanced_resume": enhanced_resume,  # ✅ Ensure enhanced_resume is included
            "cover_letter": cover_letter  # ✅ Ensure cover_letter is included
        }

    except OpenAIError as e:
        logging.error(f"OpenAI API Error: {str(e)}")
        return JSONResponse(status_code=500, content={"error": "AI service unavailable. Please try again later."})


from fastapi.responses import FileResponse

# ✅ Helper function to create PDF
def create_pdf(content, filename):
    """
    Generate a well-formatted PDF from text content.
    """
    filepath = os.path.join("generated_files", filename)
    os.makedirs("generated_files", exist_ok=True)  # Ensure directory exists

    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 50, filename.replace("_", " ").split(".")[0])

    c.setFont("Helvetica", 12)
    y_position = height - 80

    for line in content.split("\n"):
        if y_position < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = height - 50
        c.drawString(100, y_position, line)
        y_position -= 15

    c.save()
    return filepath

# ✅ Helper function to create DOCX
def create_docx(content, filename):
    """
    Generate a well-formatted DOCX file from text content.
    """
    filepath = os.path.join("generated_files", filename)
    os.makedirs("generated_files", exist_ok=True)  # Ensure directory exists

    doc = Document()
    doc.add_heading(filename.replace("_", " ").split(".")[0], level=1)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(filepath)
    return filepath

# Generate Downloadable Resumes
@app.post("/generate_resume")
async def generate_resume(data: dict):
    """
    API endpoint to generate and download a formatted resume.
    """
    resume_text = data.get("resume_text")
    file_format = data.get("format", "pdf").lower()

    if not resume_text:
        raise HTTPException(status_code=400, detail="No resume text provided")

    filename = f"Enhanced_Resume.{file_format}"
    filepath = create_pdf(resume_text, filename) if file_format == "pdf" else create_docx(resume_text, filename)

    return FileResponse(filepath, headers={"Content-Disposition": f"attachment; filename={filename}"})

# Generate Downloadable Cover Letter
@app.post("/generate_cover_letter")
async def generate_cover_letter(data: dict):
    """
    API endpoint to generate and download a formatted cover letter.
    """
    cover_letter_text = data.get("cover_letter_text")
    file_format = data.get("format", "pdf").lower()

    if not cover_letter_text:
        raise HTTPException(status_code=400, detail="No cover letter text provided")

    filename = f"Enhanced_Cover_Letter.{file_format}"
    filepath = create_pdf(cover_letter_text, filename) if file_format == "pdf" else create_docx(cover_letter_text, filename)

    return FileResponse(filepath, headers={"Content-Disposition": f"attachment; filename={filename}"})

# ✅ AI-Powered Interview with Probing Questions
# Define expected input format
class InterviewRequest(BaseModel):
    responses: Dict[int, str]
# ✅ Step 1: AI Generates Interview Questions
@app.post("/ai_interview_questions")
async def ai_interview_questions():
    """
    AI generates 5 interview questions based on the job description.
    """
    prompt = "Generate five interview questions for a candidate applying to this job."

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are an expert interviewer."},
                  {"role": "user", "content": prompt}]
    )

    # Ensure questions are formatted correctly
    questions = response.choices[0].message.content.strip().split("\n")
    structured_questions = [{"question": q} for q in questions]

    return {"questions": structured_questions}


# ✅ Step 2: AI Evaluates Answers & Provides Feedback
@app.post("/evaluate_interview")
async def evaluate_interview(request: InterviewRequest):
    """
    AI evaluates interview responses and suggests follow-up probing questions.
    """
    feedback = {}
    for index, response in request.responses.items():
        prompt = f"Evaluate this interview response and suggest a probing follow-up question: {response}"
        ai_response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "You are an expert interview coach."},
                      {"role": "user", "content": prompt}]
        )
        feedback[index] = ai_response.choices[0].message.content.strip()

    return {"interview_feedback": feedback}


# ✅ AI-Powered STAR Method Coaching
@app.post("/generate_star_response")
async def generate_star_response(scenario: str):
    prompt = f"Generate a STAR (Situation, Task, Action, Result) response for: {scenario}"
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You provide STAR-based interview answers."},
                  {"role": "user", "content": prompt}]
    )
    return {"star_response": response.choices[0].message.content.strip()}

# ✅ AI-Driven Job Matching
def is_match(skill, resume_text):
    return SequenceMatcher(None, skill.lower(), resume_text.lower()).ratio() > 0.6  # 60% match

@app.get("/match_jobs")
async def match_jobs():
    matches = []
    for seeker in job_seekers:
        for job in employer_jobs:
            if any(is_match(skill, seeker["resume_text"]) for skill in job["skills"]):
                matches.append({"job_seeker": seeker["name"], "job": job["title"]})

    return {"matches": matches}

# ✅ AI-Based Salary Benchmarking
@app.get("/salary_benchmark")
async def salary_benchmark(role: str):
    prompt = f"What is the average salary range for a {role}?"
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You provide salary insights."},
                  {"role": "user", "content": prompt}]
    )
    return {"salary_range": response.choices[0].message.content.strip()}

# ✅ AI-Based Career Path Recommendations
@app.get("/career_path")
async def career_path(role: str):
    prompt = f"What are the typical career advancement paths for a {role}?"
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You provide career guidance."},
                  {"role": "user", "content": prompt}]
    )
    return {"career_path": response.choices[0].message.content.strip()}

# ✅ AI-Based Job Offer Analysis
@app.post("/analyze_job_offer")
async def analyze_job_offer(details: str):
    prompt = f"Analyze the competitiveness of this job offer: {details}"
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You evaluate job offers."},
                  {"role": "user", "content": prompt}]
    )
    return {"job_offer_analysis": response.choices[0].message.content.strip()}

# ✅ AI-Powered Networking Recommendations
@app.get("/networking_recommendations")
async def networking_recommendations():
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You provide networking advice."},
                  {"role": "user", "content": "Suggest networking events and connections."}]
    )
    return {"networking_tips": response.choices[0].message.content.strip()}

# Interview Scheduling Automation

# ✅ AI-Suggested Interview Slots
@app.post("/schedule_interview")
async def schedule_interview(candidate_email: str, recruiter_email: str):
    prompt = (
        "Suggest three possible interview times based on availability in the next 7 days. "
        "Format the response as a list of options in the following format: \n"
        "1. Date: <YYYY-MM-DD>, Time: <HH:MM AM/PM>\n"
        "2. Date: <YYYY-MM-DD>, Time: <HH:MM AM/PM>\n"
        "3. Date: <YYYY-MM-DD>, Time: <HH:MM AM/PM>"
    )
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an AI assistant helping schedule job interviews."},
            {"role": "user", "content": prompt}
        ]
    )
    interview_slots = response.choices[0].message.content.strip().split("\n")
    
    return {"candidate_email": candidate_email, "recruiter_email": recruiter_email, "interview_slots": interview_slots}

# ✅ Confirm Interview Time
@app.post("/confirm_interview")
async def confirm_interview(candidate_email: str, recruiter_email: str, selected_time: str):
    confirmation_message = (
        f"Interview scheduled successfully!\n\n"
        f"Candidate: {candidate_email}\n"
        f"Recruiter: {recruiter_email}\n"
        f"Scheduled Time: {selected_time}"
    )
    return {"message": confirmation_message}

# FLASK APPLICATION RUNNING
#if __name__ == '__main__': - removed on advice of chatGPT - we're using FastAPI instead
 #   app.run(debug=True) - removed on advice of chatGPT
